"""Извлечение признаков таблиц.

Текущий файл — скелет. Реализацию добавляем поэтапно.
"""

from __future__ import annotations

import re

from docx import Document
from docx.oxml.ns import qn

from ...config.validation_constants import SECTION_REFERENCES
from ...config.regex_patterns import RE_TABLE_CONTINUATION, RE_TABLE_TITLE
from ...models.rich_document_structure import RunFeature, TableCellFeature, TableFeature
from .common import (
    clean_text,
    extract_run_feature,
    find_first_intro_heading_index,
    is_appendix_heading_paragraph,
    is_section_heading_paragraph,
    resolve_paragraph_alignment,
)


# Только крупные разделы после служебных блоков: иначе «РЕФЕРАТ»/«СОДЕРЖАНИЕ»
# перехватывают hint раньше «ВВЕДЕНИЯ» при обходе снизу вверх.
SECTION_HINT_MARKERS = (
    "ВВЕДЕНИЕ",
    "ЗАКЛЮЧЕНИЕ",
    SECTION_REFERENCES,
    "ПРИЛОЖЕНИЕ",
)


def _resolve_section_hint(paragraphs: list[object], anchor_paragraph_index: int | None) -> str | None:
    if anchor_paragraph_index is None:
        return None

    for paragraph in reversed(paragraphs[: anchor_paragraph_index + 1]):
        text = clean_text(paragraph.text)
        if not text:
            continue

        for marker in SECTION_HINT_MARKERS:
            if is_section_heading_paragraph(text, marker):
                return marker

    return None


def _collect_cell_runs(cell: object) -> list[RunFeature]:
    runs: list[RunFeature] = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            runs.append(extract_run_feature(run))
    return runs


def _extract_table_number(title_text: str | None) -> str | None:
    if not title_text:
        return None
    match = RE_TABLE_TITLE.search(title_text)
    if not match:
        return None
    number = match.group(2)
    return number.strip() if isinstance(number, str) and number.strip() else None


def _table_number_pattern(number_token: str | None) -> str | None:
    if not number_token:
        return None
    if number_token.isdigit():
        return "table_number_global"
    if re.match(r"^\d+\.\d+$", number_token):
        return "table_number_sectional"
    if re.match(r"^[A-Za-zА-Яа-я]\.\d+$", number_token):
        return "table_number_appendix"
    return "table_number_unknown"


def _is_valid_table_title_text(text: str | None) -> bool:
    if not text:
        return False
    return bool(RE_TABLE_TITLE.search(text) or RE_TABLE_CONTINUATION.search(text))


def _continuation_marker(title_text: str | None) -> str | None:
    if not title_text:
        return None
    if RE_TABLE_CONTINUATION.search(title_text):
        return title_text
    return None


def _collect_table_cell_map(table: object) -> list[TableCellFeature]:
    cells: list[TableCellFeature] = []
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cells.append(
                TableCellFeature(
                    row=row_index,
                    col=col_index,
                    text=clean_text(cell.text),
                    is_header_row=row_index == 0,
                    runs_features=_collect_cell_runs(cell),
                )
            )
    return cells


def _table_inside_borders(table: object) -> tuple[bool | None, bool | None]:
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        return None, None

    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        return None, None

    inside_h = tbl_borders.find(qn("w:insideH"))
    inside_v = tbl_borders.find(qn("w:insideV"))

    def _exists_border(border_element: object | None) -> bool | None:
        if border_element is None:
            return None
        val = (border_element.get(qn("w:val")) or "").lower()
        if val in {"", "nil", "none"}:
            return False
        return True

    return _exists_border(inside_h), _exists_border(inside_v)


def _table_outer_borders(table: object) -> tuple[bool | None, bool | None, bool | None, bool | None]:
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        return None, None, None, None

    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        return None, None, None, None

    top = tbl_borders.find(qn("w:top"))
    bottom = tbl_borders.find(qn("w:bottom"))
    left = tbl_borders.find(qn("w:left"))
    right = tbl_borders.find(qn("w:right"))

    def _exists_border(border_element: object | None) -> bool | None:
        if border_element is None:
            return None
        val = (border_element.get(qn("w:val")) or "").lower()
        if val in {"", "nil", "none"}:
            return False
        return True

    return _exists_border(top), _exists_border(bottom), _exists_border(left), _exists_border(right)


def _table_has_diagonal_borders(table: object) -> bool | None:
    found_cells = False
    has_diagonal = False

    for cell in table._tbl.findall(".//w:tc", table._tbl.nsmap):
        found_cells = True
        tc_pr = cell.find(qn("w:tcPr"))
        if tc_pr is None:
            continue
        tc_borders = tc_pr.find(qn("w:tcBorders"))
        if tc_borders is None:
            continue

        for diag_name in ("w:tl2br", "w:tr2bl", "w:diagUp", "w:diagDown"):
            diag = tc_borders.find(qn(diag_name))
            if diag is None:
                continue
            val = (diag.get(qn("w:val")) or "").lower()
            if val not in {"", "nil", "none"}:
                has_diagonal = True
                break

        if has_diagonal:
            break

    if not found_cells:
        return None
    return has_diagonal


def _table_prev_paragraph_map(doc: Document) -> dict[int, int | None]:
    prev_map: dict[int, int | None] = {}
    paragraph_index = 0
    table_index = 0

    for child in doc._element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph_index += 1
            continue
        if child.tag == qn("w:tbl"):
            prev_map[table_index] = paragraph_index - 1 if paragraph_index > 0 else None
            table_index += 1

    return prev_map


def _resolve_table_title_near_anchor(
    paragraphs: list[object],
    anchor_paragraph_index: int | None,
) -> tuple[str | None, int | None, str]:
    if anchor_paragraph_index is None:
        return None, None, "unknown"

    # 1) Предпочитаем стандартный случай: заголовок над таблицей.
    for probe_index in (anchor_paragraph_index, anchor_paragraph_index - 1, anchor_paragraph_index - 2):
        if not (0 <= probe_index < len(paragraphs)):
            continue

        paragraph = paragraphs[probe_index]
        candidate = clean_text(paragraph.text)
        if not candidate:
            continue

        if _is_valid_table_title_text(candidate):
            return candidate, probe_index, resolve_paragraph_alignment(paragraph)

        # Встретили непустой абзац без заголовка: дальше вверх уже не считаем заголовком этой таблицы.
        break

    # 2) Нестандартный случай: заголовок под таблицей.
    for probe_index in (anchor_paragraph_index + 1, anchor_paragraph_index + 2):
        if not (0 <= probe_index < len(paragraphs)):
            continue

        paragraph = paragraphs[probe_index]
        candidate = clean_text(paragraph.text)
        if not candidate:
            continue

        if _is_valid_table_title_text(candidate):
            return candidate, probe_index, resolve_paragraph_alignment(paragraph)

        # Непустой абзац без заголовка — глубже вниз не идем.
        break

    return None, None, "unknown"


def _table_title_relative_position(
    table_anchor_paragraph_index: int | None,
    title_paragraph_index: int | None,
) -> str:
    if table_anchor_paragraph_index is None or title_paragraph_index is None:
        return "unknown"
    if title_paragraph_index <= table_anchor_paragraph_index:
        return "above"
    if title_paragraph_index > table_anchor_paragraph_index:
        return "below"
    return "same_paragraph"


def _is_in_appendix_context(paragraphs: list[object], anchor_paragraph_index: int | None) -> bool:
    if anchor_paragraph_index is None:
        return False

    # Аналогично рисункам: после заголовка "ПРИЛОЖЕНИЕ ..." считаем, что объект в приложении.
    for paragraph in paragraphs[: anchor_paragraph_index + 1]:
        if is_appendix_heading_paragraph(clean_text(paragraph.text)):
            return True
    return False


def extract_table_features(doc: Document) -> list[TableFeature]:
    """Возвращает признаки таблиц документа.

    Этап 2 (минимум):
    - rows_count / cols_count
    - cell_text_map
    - title_above_text (эвристика по предыдущему абзацу)

    TODO:
    - Добавить анализ границ/диагоналей через XML.
    - Добавить распознавание "Продолжение таблицы N".
    """
    table_features: list[TableFeature] = []
    prev_paragraph_map = _table_prev_paragraph_map(doc)
    paragraphs = list(doc.paragraphs)
    intro_paragraph_index = find_first_intro_heading_index(paragraphs)

    for table_index, table in enumerate(doc.tables):
        rows_count = len(table.rows)
        cols_count = len(table.columns) if rows_count > 0 else 0
        prev_paragraph_index = prev_paragraph_map.get(table_index)
        if intro_paragraph_index is not None and prev_paragraph_index is not None:
            if prev_paragraph_index < intro_paragraph_index:
                continue
        if intro_paragraph_index is not None and prev_paragraph_index is None:
            continue

        title_text, title_paragraph_index, title_alignment = _resolve_table_title_near_anchor(
            paragraphs,
            prev_paragraph_index,
        )

        table_number = _extract_table_number(title_text)
        table_number_pattern = _table_number_pattern(table_number)
        title_relative_position = _table_title_relative_position(prev_paragraph_index, title_paragraph_index)
        in_appendix = _is_in_appendix_context(paragraphs, prev_paragraph_index)
        section_hint = _resolve_section_hint(paragraphs, prev_paragraph_index)

        inside_h, inside_v = _table_inside_borders(table)
        outer_top, outer_bottom, outer_left, outer_right = _table_outer_borders(table)
        cell_map = _collect_table_cell_map(table)
        table_features.append(
            TableFeature(
                table_index=table_index,
                table_anchor_paragraph_index=prev_paragraph_index,
                section_hint=section_hint,
                rows_count=rows_count,
                cols_count=cols_count,
                title_above_text=title_text,
                title_paragraph_index=title_paragraph_index,
                title_alignment=title_alignment,
                title_relative_position=title_relative_position,
                in_appendix=in_appendix,
                number=table_number,
                number_pattern=table_number_pattern,
                has_inside_horizontal_borders=inside_h,
                has_inside_vertical_borders=inside_v,
                has_outer_top_border=outer_top,
                has_outer_bottom_border=outer_bottom,
                has_outer_left_border=outer_left,
                has_outer_right_border=outer_right,
                has_diagonal_borders=_table_has_diagonal_borders(table),
                continuation_marker=_continuation_marker(title_text),
                header_row_cells=[cell for cell in cell_map if cell.is_header_row],
                first_column_cells=[cell for cell in cell_map if cell.col == 0],
                cell_text_map=cell_map,
            )
        )

    return table_features

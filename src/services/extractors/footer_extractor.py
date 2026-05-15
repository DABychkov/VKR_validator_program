"""Извлечение секционных параметров страницы и данных колонтитула.
"""

from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn

from ...models.rich_document_structure import AlignmentValue, FooterFeature, SectionPageSettingsFeature
from .common import resolve_paragraph_alignment


def _element_has_page_field(element: object) -> bool:
    for field_simple in element.findall(".//w:fldSimple", element.nsmap):
        instr = field_simple.get(qn("w:instr"), "")
        if "PAGE" in instr.upper():
            return True

    for instr_text in element.findall(".//w:instrText", element.nsmap):
        text = (instr_text.text or "").upper()
        if "PAGE" in text:
            return True

    return False


def _element_has_center_tab(element: object) -> bool:
    p_pr = element.find(qn("w:pPr"))
    if p_pr is None:
        return False

    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        return False

    for tab in tabs.findall(qn("w:tab")):
        if (tab.get(qn("w:val")) or "").lower() == "center":
            return True

    return False


def _xml_paragraph_text(paragraph_element: object) -> str:
    chunks: list[str] = []
    for text_node in paragraph_element.findall(".//w:t", paragraph_element.nsmap):
        if text_node.text:
            chunks.append(text_node.text)
    return "".join(chunks)


def _xml_paragraph_alignment(paragraph_element: object) -> AlignmentValue:
    p_pr = paragraph_element.find(qn("w:pPr"))
    if p_pr is not None:
        jc = p_pr.find(qn("w:jc"))
        if jc is not None:
            xml_value = (jc.get(qn("w:val")) or "").lower()
            if xml_value == "center":
                return "center"
            if xml_value == "left":
                return "left"
            if xml_value == "right":
                return "right"
            if xml_value in {"both", "justify"}:
                return "justify"
            if xml_value == "distribute":
                return "distribute"

    if _element_has_center_tab(paragraph_element):
        return "center"

    return "unknown"


def _xml_explicit_alignment(paragraph_element: object) -> AlignmentValue | None:
    """Возвращает явно заданный w:jc (без эвристик)."""
    p_pr = paragraph_element.find(qn("w:pPr"))
    if p_pr is None:
        return None

    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        return None

    xml_value = (jc.get(qn("w:val")) or "").lower()
    if xml_value == "center":
        return "center"
    if xml_value == "left":
        return "left"
    if xml_value == "right":
        return "right"
    if xml_value in {"both", "justify"}:
        return "justify"
    if xml_value == "distribute":
        return "distribute"
    return "unknown"


def _paragraph_has_page_field(paragraph: object) -> bool:
    return _element_has_page_field(paragraph._element)


def _paragraph_has_center_tab(paragraph: object) -> bool:
    return _element_has_center_tab(paragraph._element)


def _resolve_footer_paragraph_alignment(paragraph: object) -> AlignmentValue:
    alignment = resolve_paragraph_alignment(paragraph, default_left=False)
    if alignment != "unknown":
        return alignment

    # Часто PAGE визуально центрируют через таб-стоп, а не через paragraph jc.
    if _paragraph_has_center_tab(paragraph):
        return "center"

    return "unknown"


def _footer_alignment(section: object) -> AlignmentValue:
    footer = section.footer
    footer_element = footer._element

    # Сначала обходим все XML-абзацы футера, включая w:sdt/w:sdtContent,
    # которые python-docx не всегда поднимает в footer.paragraphs.
    for paragraph_element in footer_element.findall(".//w:p", footer_element.nsmap):
        if not _element_has_page_field(paragraph_element):
            continue

        explicit_alignment = _xml_explicit_alignment(paragraph_element)
        if explicit_alignment in {"left", "right", "justify", "distribute"}:
            return explicit_alignment
        if explicit_alignment == "center":
            return "center"
        if _element_has_center_tab(paragraph_element):
            return "center"

        # Для чистого поля PAGE без явного выравнивания считаем центр,
        # т.к. в шаблонах часто центр задается на уровне стиля/секции.
        if not _xml_paragraph_text(paragraph_element).strip():
            return "center"
        return "center"

    # Приоритет: абзац, где реально стоит поле PAGE.
    for paragraph in footer.paragraphs:
        if _paragraph_has_page_field(paragraph):
            alignment = _resolve_footer_paragraph_alignment(paragraph)
            if alignment != "unknown":
                return alignment

            # Практический fallback: если это чистый PAGE-параграф,
            # в типовых шаблонах нумерация обычно центрирована.
            if not paragraph.text.strip():
                return "center"

            return alignment

    for paragraph in footer.paragraphs:
        if paragraph.text.strip():
            return _resolve_footer_paragraph_alignment(paragraph)

    if footer.paragraphs:
        return _resolve_footer_paragraph_alignment(footer.paragraphs[0])
    return "unknown"


def _has_page_field(section: object) -> bool:
    return _element_has_page_field(section.footer._element)


def _extract_pg_num_type(section: object) -> tuple[bool | None, int | None, str | None]:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        return None, None, None

    start_value = pg_num_type.get(qn("w:start"))
    restart_numbering = start_value is not None
    start_number = int(start_value) if start_value and start_value.isdigit() else None
    number_format = pg_num_type.get(qn("w:fmt"))
    return restart_numbering, start_number, number_format


def _start_type_name(section: object) -> str | None:
    start_type = section.start_type
    if start_type is None:
        return None
    return getattr(start_type, "name", str(start_type))


def extract_section_page_settings(doc: Document) -> list[SectionPageSettingsFeature]:
    """Возвращает параметры страницы по секциям (поля/размер/ориентация)."""
    settings: list[SectionPageSettingsFeature] = []

    for section_index, section in enumerate(doc.sections):
        settings.append(
            SectionPageSettingsFeature(
                section_index=section_index,
                start_type=_start_type_name(section),
                page_width_mm=float(section.page_width.mm) if section.page_width is not None else None,
                page_height_mm=float(section.page_height.mm) if section.page_height is not None else None,
                margin_left_mm=float(section.left_margin.mm) if section.left_margin is not None else None,
                margin_right_mm=float(section.right_margin.mm) if section.right_margin is not None else None,
                margin_top_mm=float(section.top_margin.mm) if section.top_margin is not None else None,
                margin_bottom_mm=float(section.bottom_margin.mm)
                if section.bottom_margin is not None
                else None,
            )
        )

    return settings


def extract_footer_features(doc: Document) -> list[FooterFeature]:
    """Возвращает признаки пагинации из футеров.
    """
    footer_features: list[FooterFeature] = []

    for section_index, section in enumerate(doc.sections):
        restart_numbering, start_number, number_format = _extract_pg_num_type(section)

        footer_features.append(
            FooterFeature(
                section_index=section_index,
                has_page_field=_has_page_field(section),
                page_field_format=number_format,
                alignment=_footer_alignment(section),
                restart_numbering=restart_numbering,
                start_number=start_number,
            )
        )

    return footer_features

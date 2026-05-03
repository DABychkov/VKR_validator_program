"""Извлечение признаков примечаний и сносок."""

from __future__ import annotations

from lxml import etree
from docx import Document
from docx.oxml.ns import qn

from ...config.regex_patterns import (
    RE_ASTERISK_FOOTNOTE_BODY,
    RE_ASTERISK_FOOTNOTE_INLINE_MARKER,
    RE_FIGURE_CAPTION,
    RE_NOTES_HEADER,
    RE_NOTES_ITEM,
    RE_NOTES_ITEM_NUMBER_PREFIX,
    RE_NOTE_SINGLE,
    RE_NUMBERED_NOTE_SINGLE,
    RE_TABLE_CONTINUATION,
    RE_TABLE_TITLE,
)
from ...models.rich_document_structure import FootnoteFeature, NoteFeature
from .common import clean_text


def _has_dot_after_number(text: str) -> bool | None:
    """Возвращает флаг наличия точки после номера в начале строки примечания."""
    match = RE_NOTES_ITEM_NUMBER_PREFIX.match(str(text or ""))
    if not match:
        return None
    return match.group(2) == "."


def _is_note_word_prefix(text: str) -> bool:
    lowered = str(text or "").lower().strip()
    return lowered.startswith("примечание") or lowered.startswith("примечания")


def _has_numbered_item_in_next_non_empty_window(
    paragraphs: list[str],
    start_index: int,
    *,
    max_non_empty_lines: int = 3,
) -> bool:
    non_empty_seen = 0
    for probe_index in range(start_index + 1, len(paragraphs)):
        candidate = str(paragraphs[probe_index] or "").strip()
        if not candidate:
            continue

        non_empty_seen += 1
        if RE_NOTES_ITEM.match(candidate):
            return True
        if non_empty_seen >= max_non_empty_lines:
            return False

    return False


def _extract_single_note_body_without_dash(text: str) -> str:
    lowered = str(text or "").lower().strip()
    if not lowered.startswith("примечание"):
        return ""

    body = str(text or "").strip()[len("Примечание") :].lstrip(" :\t")
    return body.strip()


def _near_material_flags(
    paragraph_index: int,
    figure_caption_indices: set[int],
    table_caption_indices: set[int],
    window: int = 5,
) -> tuple[bool, bool]:
    near_figure = any(abs(paragraph_index - idx) <= window for idx in figure_caption_indices)
    near_table = any(abs(paragraph_index - idx) <= window for idx in table_caption_indices)
    return near_figure, near_table


def _footnotes_root(doc: Document) -> object | None:
    for part in doc.part.package.parts:
        part_name = str(getattr(part, "partname", ""))
        if not part_name.endswith("/footnotes.xml"):
            continue

        root = getattr(part, "_element", None)
        if root is None:
            blob = getattr(part, "blob", None)
            if not blob:
                return None
            root = etree.fromstring(blob)
        return root

    return None


def _extract_footnote_ids_from_part(doc: Document) -> set[int]:
    """Возвращает набор id из /word/footnotes.xml, если часть присутствует."""
    root = _footnotes_root(doc)
    if root is None:
        return set()

    result: set[int] = set()
    for node in root.findall(".//w:footnote", root.nsmap):
        raw_id = node.get(qn("w:id"))
        if raw_id is None:
            continue
        try:
            result.add(int(raw_id))
        except ValueError:
            continue
    return result


def _extract_footnote_marker_from_content(doc: Document, footnote_id: int) -> str | None:
    """
    Читает первый символ/слово из содержимого сноски в footnotes.xml.
    Возвращает '*' если сноска начинается с *, иначе None (используем ID как маркер).
    """
    root = _footnotes_root(doc)
    if root is None:
        return None

    for node in root.findall(".//w:footnote", root.nsmap):
        raw_id = node.get(qn("w:id"))
        if raw_id is None:
            continue
        try:
            if int(raw_id) != footnote_id:
                continue
        except ValueError:
            continue

        # Читаем текст сноски
        text_parts = []
        for t_elem in node.findall(".//w:t", root.nsmap):
            if t_elem.text:
                text_parts.append(t_elem.text)

        footnote_text = "".join(text_parts).strip()
        if footnote_text.startswith("*"):
            return "*"
        break

    return None


def _extract_footnote_separator_flags(doc: Document) -> tuple[bool | None, bool | None]:
    """Возвращает признаки разделительной линии для сносок из footnotes.xml."""
    root = _footnotes_root(doc)
    if root is None:
        return None, None

    separator_node = None
    for node in root.findall(".//w:footnote", root.nsmap):
        node_type = (node.get(qn("w:type")) or "").lower()
        if node_type == "separator":
            separator_node = node
            break

    if separator_node is None:
        return False, None

    has_separator = bool(
        separator_node.findall(".//w:separator", separator_node.nsmap)
        or separator_node.findall(".//w:pBdr/w:top", separator_node.nsmap)
    )
    # Для встроенной xml-сноски опираемся на стандартный Word separator.
    short_left_heuristic = has_separator
    return has_separator, short_left_heuristic


def _paragraph_has_visual_separator_line(paragraph: object) -> bool:
    p = getattr(paragraph, "_p", None)
    if p is None:
        return False
    text = str(getattr(paragraph, "text", "") or "").strip()

    # Граница абзаца (горизонтальная линия через borders).
    has_border_line = bool(
        p.findall(".//w:pPr/w:pBdr/w:top", p.nsmap)
        or p.findall(".//w:pPr/w:pBdr/w:bottom", p.nsmap)
    )
    if has_border_line:
        return True

    # Word иногда вставляет горизонтальную линию как drawing/pict объект.
    has_drawing_line = bool(
        p.findall(".//w:drawing", p.nsmap)
        or p.findall(".//w:pict", p.nsmap)
    )
    # Для drawing/pict считаем линией только отдельный пустой абзац,
    # чтобы не ловить обычные встроенные объекты в текстовых строках.
    return has_drawing_line and not text


def _parse_width_pt_from_style(style: str | None) -> float | None:
    if not style:
        return None

    for token in str(style).split(";"):
        normalized = token.strip().lower().replace(" ", "")
        if not normalized.startswith("width:"):
            continue
        raw_value = normalized[len("width:") :]
        if not raw_value.endswith("pt"):
            continue
        numeric_part = raw_value[:-2]
        try:
            return float(numeric_part)
        except ValueError:
            return None

    return None


def _text_width_pt(doc: Document) -> float | None:
    try:
        section = doc.sections[0]
    except Exception:
        return None

    page_width = getattr(section, "page_width", None)
    left_margin = getattr(section, "left_margin", None)
    right_margin = getattr(section, "right_margin", None)
    if page_width is None or left_margin is None or right_margin is None:
        return None

    width_pt = float(page_width.pt - left_margin.pt - right_margin.pt)
    if width_pt <= 0:
        return None
    return width_pt


def _custom_separator_short_heuristic(paragraph: object, doc: Document) -> bool | None:
    p = getattr(paragraph, "_p", None)
    if p is None:
        return None

    # 1) VML horizontal rule percent (наиболее надежный для вставленной Word линии).
    office_ns = "{urn:schemas-microsoft-com:office:office}"
    for rect in p.findall(".//v:rect", p.nsmap):
        if (rect.get(f"{office_ns}hr") or "").lower() != "t":
            continue

        hrpct_raw = rect.get(f"{office_ns}hrpct")
        if hrpct_raw is not None:
            try:
                hrpct_value = float(hrpct_raw)
                # o:hrpct обычно хранится в десятых долях процента (405 => 40.5%).
                ratio = hrpct_value / 10.0 if hrpct_value > 100 else hrpct_value
                return ratio <= 60.0
            except ValueError:
                pass

        width_pt = _parse_width_pt_from_style(rect.get("style"))
        text_width = _text_width_pt(doc)
        if width_pt is not None and text_width is not None:
            return width_pt <= text_width * 0.6

    # 2) Альтернативно пробуем drawing extent (если будет OOXML drawing линия).
    for extent in p.findall(".//wp:extent", p.nsmap):
        cx_raw = extent.get("cx")
        if cx_raw is None:
            continue
        try:
            width_emu = float(cx_raw)
        except ValueError:
            continue

        # 1 pt = 12700 EMU.
        width_pt = width_emu / 12700.0
        text_width = _text_width_pt(doc)
        if text_width is not None:
            return width_pt <= text_width * 0.6

    return None


def _extract_custom_asterisk_separator_flags(doc: Document, paragraphs: list[object]) -> tuple[bool | None, bool | None]:
    """Пытается определить наличие/короткость линии для кастомных звездочных сносок."""
    body_indices = [
        index
        for index, paragraph in enumerate(paragraphs)
        if RE_ASTERISK_FOOTNOTE_BODY.match(getattr(paragraph, "text", "") or "")
    ]
    if not body_indices:
        return None, None

    line_found = False
    candidate_line_paragraphs: list[object] = []
    for body_index in body_indices:
        # Ищем линию перед текстом сноски: обычно это предыдущий абзац(ы).
        start = max(0, body_index - 2)
        end = min(len(paragraphs), body_index + 1)
        for probe_index in range(start, end):
            if _paragraph_has_visual_separator_line(paragraphs[probe_index]):
                line_found = True
                candidate_line_paragraphs.append(paragraphs[probe_index])
                break
        if line_found:
            break

    if not line_found:
        return False, None

    for paragraph in candidate_line_paragraphs:
        short_value = _custom_separator_short_heuristic(paragraph, doc)
        if short_value is not None:
            return True, short_value

    # Если линия обнаружена, но длину вычислить нельзя, оставляем heuristic=None.
    return True, None


def extract_notes_features(doc: Document) -> list[NoteFeature]:
    """Извлекает абзацы-примечания по шаблонам ТЗ/ГОСТ."""
    notes: list[NoteFeature] = []
    in_notes_group = False

    normalized_paragraphs = [clean_text(paragraph.text) for paragraph in doc.paragraphs]
    figure_caption_indices = {
        index for index, text in enumerate(normalized_paragraphs) if text and RE_FIGURE_CAPTION.match(text)
    }
    table_caption_indices = {
        index
        for index, text in enumerate(normalized_paragraphs)
        if text and (RE_TABLE_TITLE.match(text) or RE_TABLE_CONTINUATION.search(text))
    }

    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        text = clean_text(paragraph.text)
        if not text:
            in_notes_group = False
            continue

        near_figure, near_table = _near_material_flags(paragraph_index, figure_caption_indices, table_caption_indices)
        near_material = near_figure or near_table

        if RE_NOTE_SINGLE.match(text):
            notes.append(
                NoteFeature(
                    paragraph_index=paragraph_index,
                    raw_text=text,
                    note_kind="single",
                    has_dash_separator=True,
                    near_figure_caption=near_figure,
                    near_table_caption=near_table,
                )
            )
            in_notes_group = False
            continue

        numbered_single_match = RE_NUMBERED_NOTE_SINGLE.match(text)
        if numbered_single_match:
            notes.append(
                NoteFeature(
                    paragraph_index=paragraph_index,
                    raw_text=text,
                    note_kind="numbered_single",
                    item_number=int(numbered_single_match.group(1)),
                    has_dot_after_number=_has_dot_after_number(text),
                    has_dash_separator=True,
                    near_figure_caption=near_figure,
                    near_table_caption=near_table,
                )
            )
            in_notes_group = False
            continue

        if RE_NOTES_HEADER.match(text):
            has_numbered_after = _has_numbered_item_in_next_non_empty_window(
                normalized_paragraphs,
                paragraph_index,
                max_non_empty_lines=3,
            )
            if has_numbered_after:
                notes.append(
                    NoteFeature(
                        paragraph_index=paragraph_index,
                        raw_text=text,
                        note_kind="group_header",
                        has_dash_separator=False,
                        near_figure_caption=near_figure,
                        near_table_caption=near_table,
                    )
                )
                in_notes_group = True
                continue

        # fallback для "похожих на примечание" строк:
        # применяем только рядом с таблицей/рисунком, чтобы не ловить обычный текст.
        if near_material and _is_note_word_prefix(text):
            has_numbered_after = _has_numbered_item_in_next_non_empty_window(
                normalized_paragraphs,
                paragraph_index,
                max_non_empty_lines=3,
            )
            if has_numbered_after:
                notes.append(
                    NoteFeature(
                        paragraph_index=paragraph_index,
                        raw_text=text,
                        note_kind="group_header",
                        has_dash_separator=False,
                        near_figure_caption=near_figure,
                        near_table_caption=near_table,
                    )
                )
                in_notes_group = True
                continue

            body_without_dash = _extract_single_note_body_without_dash(text)
            if body_without_dash:
                notes.append(
                    NoteFeature(
                        paragraph_index=paragraph_index,
                        raw_text=text,
                        note_kind="single",
                        has_dash_separator=False,
                        near_figure_caption=near_figure,
                        near_table_caption=near_table,
                    )
                )
                in_notes_group = False
                continue

            notes.append(
                NoteFeature(
                    paragraph_index=paragraph_index,
                    raw_text=text,
                    note_kind="unknown",
                    has_dash_separator=None,
                    near_figure_caption=near_figure,
                    near_table_caption=near_table,
                )
            )
            in_notes_group = False
            continue

        if in_notes_group:
            group_item_match = RE_NOTES_ITEM.match(text)
            if group_item_match:
                notes.append(
                    NoteFeature(
                        paragraph_index=paragraph_index,
                        raw_text=text,
                        note_kind="group_item",
                        item_number=int(group_item_match.group(1)),
                        has_dot_after_number=_has_dot_after_number(text),
                        has_dash_separator=None,
                        near_figure_caption=near_figure,
                        near_table_caption=near_table,
                    )
                )
                continue
            in_notes_group = False

    return notes


def extract_footnote_features(doc: Document) -> list[FootnoteFeature]:
    """Извлекает маркеры сносок из body и связывает с /word/footnotes.xml."""
    footnote_features: list[FootnoteFeature] = []
    existing_footnote_ids = _extract_footnote_ids_from_part(doc)
    xml_has_separator_line, xml_separator_short_left = _extract_footnote_separator_flags(doc)

    paragraphs = list(doc.paragraphs)
    custom_has_separator_line, custom_separator_short_left = _extract_custom_asterisk_separator_flags(doc, paragraphs)
    asterisk_body_indices = [
        index
        for index, paragraph in enumerate(paragraphs)
        if RE_ASTERISK_FOOTNOTE_BODY.match(paragraph.text or "")
    ]
    inline_asterisk_markers: list[int] = []

    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text or ""
        inline_asterisk_markers.extend([index for _ in RE_ASTERISK_FOOTNOTE_INLINE_MARKER.finditer(text)])

    asterisk_body_cursor = 0

    for paragraph_index, paragraph in enumerate(paragraphs):
        paragraph_element = paragraph._p

        for reference in paragraph_element.findall(".//w:footnoteReference", paragraph_element.nsmap):
            raw_id = reference.get(qn("w:id"))
            if raw_id is None:
                continue

            footnote_id: int | None = None
            try:
                footnote_id = int(raw_id)
            except ValueError:
                pass

            # Читаем маркер из содержимого сноски (если он начинается с *)
            custom_marker = _extract_footnote_marker_from_content(doc, footnote_id) if footnote_id is not None else None
            is_custom_marker = custom_marker == "*"
            marker_text = "*" if is_custom_marker else str(raw_id)

            footnote_features.append(
                FootnoteFeature(
                    paragraph_index=paragraph_index,
                    marker_text=marker_text,
                    marker_type="xml_reference",
                    footnote_id=footnote_id,
                    custom_mark_follows=is_custom_marker,
                    resolved_in_footnotes_part=(footnote_id in existing_footnote_ids) if footnote_id is not None else None,
                    has_separator_line=xml_has_separator_line,
                    separator_short_left_heuristic=xml_separator_short_left,
                )
            )

        text = paragraph.text or ""
        for _ in RE_ASTERISK_FOOTNOTE_INLINE_MARKER.finditer(text):
            # Если уже есть xml_reference с marker=*, не добавляем asterisk для этого абзаца
            has_xml_reference_star = any(
                f.paragraph_index == paragraph_index and f.marker_type == "xml_reference" and f.marker_text == "*"
                for f in footnote_features
            )
            if has_xml_reference_star:
                continue

            # Резолвим каждый звездочный маркер в ближайшее тело сноски ниже по документу.
            while (
                asterisk_body_cursor < len(asterisk_body_indices)
                and asterisk_body_indices[asterisk_body_cursor] <= paragraph_index
            ):
                asterisk_body_cursor += 1

            has_asterisk_resolution = asterisk_body_cursor < len(asterisk_body_indices)
            if has_asterisk_resolution:
                asterisk_body_cursor += 1

            footnote_features.append(
                FootnoteFeature(
                    paragraph_index=paragraph_index,
                    marker_text="*",
                    marker_type="asterisk",
                    footnote_id=None,
                    custom_mark_follows=None,
                    resolved_in_footnotes_part=has_asterisk_resolution,
                    has_separator_line=custom_has_separator_line,
                    separator_short_left_heuristic=custom_separator_short_left,
                )
            )

    return footnote_features

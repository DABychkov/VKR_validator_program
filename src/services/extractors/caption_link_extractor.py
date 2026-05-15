"""Извлечение подписей рисунков и ссылок из текста документа."""

from __future__ import annotations

import re

from docx import Document
from docx.text.paragraph import Paragraph

from ...config.regex_patterns import (
    RE_FIGURE_CAPTION,
    RE_FIGURE_LINK,
    RE_FORMULA_LINK,
    RE_REFERENCE_LIST_ITEM,
    RE_SOURCE_LINK,
    RE_TABLE_CONTINUATION,
    RE_TABLE_LINK,
    RE_TABLE_TITLE,
)
from ...models.rich_document_structure import FigureCaptionFeature, FormulaFeature, LinkFeature, TableFeature
from .common import find_first_intro_heading_index, is_appendix_heading_paragraph, resolve_paragraph_alignment


def _is_range(target: str | None) -> bool:
    if not target:
        return False
    return bool(re.search(r"[-\u2013]", target))


def _extract_references_numbers(doc: Document) -> set[int] | None:
    """Возвращает номера источников из раздела "Список использованных источников".

    Если раздел не найден, возвращает None.
    """
    paragraphs = [" ".join(p.text.split()) for p in doc.paragraphs]

    start_index = None
    for idx, text in enumerate(paragraphs):
        upper = text.upper()
        if "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ" in upper:
            start_index = idx
            break

    if start_index is None:
        return None

    numbers: set[int] = set()

    for text in paragraphs[start_index + 1 :]:
        if not text:
            continue

        # Доходим до приложений и выходим из секции источников.
        if is_appendix_heading_paragraph(text):
            break

        match = RE_REFERENCE_LIST_ITEM.match(text)
        if match:
            numbers.add(int(match.group(1)))

    return numbers


def _resolve_source_link(start_number: str, end_number: str | None, refs: set[int] | None) -> bool | None:
    if refs is None:
        return None

    try:
        start = int(start_number)
    except ValueError:
        return False

    if end_number is None:
        return start in refs

    try:
        end = int(end_number)
    except ValueError:
        return False

    if end < start:
        return False

    return all(num in refs for num in range(start, end + 1))


def _figure_caption_pattern(number_token: str) -> str:
    if re.match(r"^[A-Za-zА-Яа-я]\.\d+$", number_token):
        return "figure_caption_appendix"
    if re.match(r"^\d+(?:\.\d+)+$", number_token):
        return "figure_caption_sectional"
    if re.match(r"^\d+$", number_token):
        return "figure_caption_global"
    return "figure_caption_unknown"


def _normalize_target_number(token: str | None) -> str | None:
    if token is None:
        return None
    cleaned = re.sub(r"[\s\(\)]", "", token)
    return cleaned or None


def _extract_figure_number(text: str) -> str | None:
    match = RE_FIGURE_CAPTION.search(text)
    if not match:
        return None
    return _normalize_target_number(match.group(2))


def _extract_table_number(text: str | None) -> str | None:
    return _normalize_target_number(text)


def _extract_formula_number(number_value: str | None) -> str | None:
    if not number_value:
        return None
    return _normalize_target_number(number_value)


def _paragraph_has_drawing(paragraph: Paragraph) -> bool:
    paragraph_element = paragraph._p
    drawing_count = len(paragraph_element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"))
    pict_count = len(paragraph_element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict"))
    return (drawing_count + pict_count) > 0


def _has_nearby_drawing(paragraphs: list[Paragraph], paragraph_index: int) -> bool:
    # В реальных DOCX подпись может идти сразу под рисунком, поэтому проверяем окно вокруг подписи.
    for offset in (-2, -1, 0, 1):
        probe_index = paragraph_index + offset
        if 0 <= probe_index < len(paragraphs) and _paragraph_has_drawing(paragraphs[probe_index]):
            return True
    return False


def _nearest_drawing_offset(paragraphs: list[Paragraph], paragraph_index: int) -> int | None:
    nearest_offset: int | None = None
    for offset in (-2, -1, 0, 1):
        probe_index = paragraph_index + offset
        if not (0 <= probe_index < len(paragraphs)):
            continue
        if not _paragraph_has_drawing(paragraphs[probe_index]):
            continue
        if nearest_offset is None or abs(offset) < abs(nearest_offset):
            nearest_offset = offset
    return nearest_offset


def _drawing_relative_position(paragraphs: list[Paragraph], paragraph_index: int) -> str:
    nearest_offset = _nearest_drawing_offset(paragraphs, paragraph_index)
    if nearest_offset is None:
        return "unknown"
    if nearest_offset < 0:
        return "below"
    if nearest_offset > 0:
        return "above"
    return "same_paragraph"


def extract_figure_caption_features(doc: Document) -> list[FigureCaptionFeature]:
    """Извлекает подписи рисунков по regex-эвристике."""
    result: list[FigureCaptionFeature] = []
    is_in_appendix_context = False

    paragraphs = list(doc.paragraphs)
    intro_paragraph_index = find_first_intro_heading_index(paragraphs)
    for paragraph_index, paragraph in enumerate(paragraphs):
        text = " ".join(paragraph.text.split())
        if not text:
            continue

        # Якорь приложения: после заголовка «ПРИЛОЖЕНИЕ …» (также с хвостом строки).
        if (
            is_appendix_heading_paragraph(text)
            and (intro_paragraph_index is None or paragraph_index >= intro_paragraph_index)
        ):
            is_in_appendix_context = True

        match = RE_FIGURE_CAPTION.search(text)
        if not match:
            continue

        number_token = match.group(2)
        separator = match.group(3) or ""
        result.append(
            FigureCaptionFeature(
                paragraph_index=paragraph_index,
                caption_text=text,
                caption_number=_extract_figure_number(text),
                alignment=resolve_paragraph_alignment(paragraph),
                pattern_type=_figure_caption_pattern(number_token),
                has_dash_separator=separator in {"-", "–", "—", ":"},
                ends_with_period=text.endswith("."),
                has_nearby_drawing=_has_nearby_drawing(paragraphs, paragraph_index),
                drawing_relative_position=_drawing_relative_position(paragraphs, paragraph_index),
                in_appendix=is_in_appendix_context,
            )
        )

    return result


def extract_links_features(doc: Document) -> list[LinkFeature]:
    """Извлекает ссылки на источники/рисунки/таблицы/формулы."""
    links: list[LinkFeature] = []
    refs_numbers = _extract_references_numbers(doc)

    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        text = " ".join(paragraph.text.split())
        if not text:
            continue

        is_caption_or_title = bool(
            RE_FIGURE_CAPTION.match(text)
            or RE_TABLE_TITLE.match(text)
            or RE_TABLE_CONTINUATION.search(text)
        )

        for match in RE_SOURCE_LINK.finditer(text):
            start_number = match.group("start")
            end_number = match.group("end")
            target = f"{start_number}-{end_number}" if end_number else start_number
            links.append(
                LinkFeature(
                    link_type="source",
                    paragraph_index=paragraph_index,
                    raw_text=match.group(0),
                    target_number=target,
                    is_range=end_number is not None or _is_range(target),
                    resolved_in_target_list=_resolve_source_link(start_number, end_number, refs_numbers),
                )
            )

        if is_caption_or_title:
            continue

        for match in RE_FIGURE_LINK.finditer(text):
            target = match.group(1)
            links.append(
                LinkFeature(
                    link_type="figure",
                    paragraph_index=paragraph_index,
                    raw_text=match.group(0),
                    target_number=target,
                )
            )

        for match in RE_TABLE_LINK.finditer(text):
            raw = match.group(0)
            target = _normalize_target_number(match.group(1))
            links.append(
                LinkFeature(
                    link_type="table",
                    paragraph_index=paragraph_index,
                    raw_text=raw,
                    target_number=target,
                )
            )

        for match in RE_FORMULA_LINK.finditer(text):
            raw = match.group(0)
            target = _normalize_target_number(match.group(1))
            links.append(
                LinkFeature(
                    link_type="formula",
                    paragraph_index=paragraph_index,
                    raw_text=raw,
                    target_number=target,
                )
            )

    return links


def resolve_non_source_links(
    links: list[LinkFeature],
    figure_caption_features: list[FigureCaptionFeature],
    table_features: list[TableFeature],
    formula_features: list[FormulaFeature],
) -> list[LinkFeature]:
    """Проставляет resolved-флаги для figure/table/formula ссылок."""

    def _set_resolution(
        link: LinkFeature,
        in_list: bool | None,
        with_object: bool | None,
        relative: str | None = None,
    ) -> None:
        link.resolved_in_target_list = in_list
        link.resolved_with_object = with_object
        link.relative_position_to_target = relative

    def _relative_position(link_paragraph_index: int | None, target_paragraph_index: int | None) -> str:
        if link_paragraph_index is None or target_paragraph_index is None:
            return "unknown"
        if link_paragraph_index < target_paragraph_index:
            return "before"
        if link_paragraph_index > target_paragraph_index:
            return "after"
        return "same"

    figure_by_number: dict[str, list[FigureCaptionFeature]] = {}
    for caption in figure_caption_features:
        number = _normalize_target_number(caption.caption_number)
        if not number:
            continue
        figure_by_number.setdefault(number, []).append(caption)

    table_index_by_number: dict[str, int] = {}
    table_numbers: set[str] = set()
    for table in table_features:
        number = _extract_table_number(getattr(table, "number", None))
        if not number:
            continue
        table_numbers.add(number)
        anchor_index = getattr(table, "table_anchor_paragraph_index", None)
        title_index = getattr(table, "title_paragraph_index", None)
        target_index = anchor_index if anchor_index is not None else title_index
        if target_index is not None:
            table_index_by_number.setdefault(number, int(target_index))

    formula_by_number: dict[str, list[FormulaFeature]] = {}
    for formula in formula_features:
        number = _extract_formula_number(getattr(formula, "number", None))
        if not number:
            continue
        formula_by_number.setdefault(number, []).append(formula)

    for link in links:
        if link.link_type == "source":
            in_list = link.resolved_in_target_list
            with_object = True if in_list is True else (False if in_list is False else None)
            _set_resolution(link, in_list, with_object, "unknown")
            continue

        target_number = _normalize_target_number(link.target_number)
        if not target_number:
            _set_resolution(link, False, None, "unknown")
            continue

        if link.link_type == "figure":
            matched_captions = figure_by_number.get(target_number, [])
            target_index = min((c.paragraph_index for c in matched_captions), default=None)
            _set_resolution(
                link,
                bool(matched_captions),
                any(caption.has_nearby_drawing for caption in matched_captions),
                _relative_position(link.paragraph_index, target_index),
            )
            continue

        if link.link_type == "table":
            exists = target_number in table_numbers
            # Для таблиц "object" = реально найденная таблица с указанным номером.
            _set_resolution(
                link,
                exists,
                exists,
                _relative_position(link.paragraph_index, table_index_by_number.get(target_number)),
            )
            continue

        if link.link_type == "formula":
            matched_formulas = formula_by_number.get(target_number, [])
            target_index = min((f.paragraph_index for f in matched_formulas), default=None)
            _set_resolution(
                link,
                bool(matched_formulas),
                any(f.omml_xml is not None for f in matched_formulas),
                _relative_position(link.paragraph_index, target_index),
            )
            continue

        _set_resolution(link, None, None, "unknown")

    return links
